# utils/auth.py
from passlib.context import CryptContext

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def hash_password(password: str) -> str:
    """Hash a plaintext password."""
    return pwd_context.hash(password)

def verify_password(password: str, hashed: str) -> bool:
    """Verify a plaintext password against the hashed version."""
    return pwd_context.verify(password, hashed)


# utils/db.py
import os
from sqlalchemy import create_engine, Column, Integer, String, Text, Float, JSON, ForeignKey
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, relationship

DB_URL = os.getenv("DATABASE_URL", "sqlite:///ketos_resumes.db")
engine = create_engine(DB_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True)
    name = Column(String)
    password = Column(String)

class Job(Base):
    __tablename__ = "jobs"
    id = Column(Integer, primary_key=True, index=True)
    title = Column(String)
    description = Column(Text)
    criteria = Column(JSON)
    threshold = Column(Float, default=0.0)
    candidates = relationship("Candidate", back_populates="job")

class Candidate(Base):
    __tablename__ = "candidates"
    id = Column(Integer, primary_key=True, index=True)
    job_id = Column(Integer, ForeignKey("jobs.id"))
    name = Column(String)
    email = Column(String)
    phone = Column(String)
    score = Column(Float)
    summary = Column(Text)
    status = Column(String, default="Pending")
    raw_data = Column(JSON)
    job = relationship("Job", back_populates="candidates")

# Utility functions
def get_db():
    Base.metadata.create_all(bind=engine)
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def get_current_job(db, job_id=None):
    # For simplicity, return most recent job if no ID provided
    return db.query(Job).order_by(Job.id.desc()).first()


# utils/normalization.py
from dateutil import parser

SKILL_MAPPING = {
    "ms office": "Microsoft Office",
    "excel": "Microsoft Excel",
    # Add more mappings as needed
}

def normalize_date(date_str: str) -> str:
    """Convert various date formats to YYYY-MM"""
    try:
        dt = parser.parse(date_str)
        return dt.strftime("%Y-%m")
    except Exception:
        return date_str


def normalize_skill(skill: str) -> str:
    key = skill.strip().lower()
    return SKILL_MAPPING.get(key, skill.strip())


# utils/parser.py
import io
import re
from pdfplumber import open as open_pdf
from docx import Document
from utils.normalization import normalize_date, normalize_skill

EMAIL_RE = re.compile(r"[a-zA-Z0-9+._%-]+@[a-zA-Z0-9._%-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"\+?\d[\d \-]{7,}\d")


def parse_job_description(file_obj) -> str:
    # Simplest: extract full text
    if hasattr(file_obj, "read") and file_obj.name.lower().endswith('.pdf'):
        text = ''
        with open_pdf(file_obj) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ''
        return text
    elif hasattr(file_obj, "read") and file_obj.name.lower().endswith('.docx'):
        doc = Document(io.BytesIO(file_obj.read()))
        return '\n'.join(p.text for p in doc.paragraphs)
    return ''


def parse_criteria(file_obj) -> str:
    # reuse same logic as JD
    return parse_job_description(file_obj)


def parse_resume(file_obj) -> dict:
    text = ''
    if file_obj.name.lower().endswith('.pdf'):
        with open_pdf(file_obj) as pdf:
            for pg in pdf.pages:
                text += pg.extract_text() or ''
    elif file_obj.name.lower().endswith('.docx'):
        doc = Document(io.BytesIO(file_obj.read()))
        text = '\n'.join(p.text for p in doc.paragraphs)

    # Extract basic fields
    name = text.split('\n')[0].strip()
    email = EMAIL_RE.search(text)
    phone = PHONE_RE.search(text)

    # Skills: simple keyword scan (comma or newline separated)
    skills = []
    for line in text.splitlines():
        if 'skills' in line.lower():
            parts = re.split(r'[:,]', line, maxsplit=1)
            if len(parts) > 1:
                skills = [normalize_skill(s) for s in parts[1].split(',')]
                break

    # Dates normalization (simplest, find patterns)
    dates = re.findall(r'\b\w+ \d{4}\b', text)
    dates = [normalize_date(d) for d in dates]

    return {
        'name': name,
        'email': email.group(0) if email else None,
        'phone': phone.group(0) if phone else None,
        'skills': skills,
        'dates': dates,
        'raw_text': text
    }


# utils/screening.py
from utils.normalization import normalize_skill

def score_candidate(fields: dict, criteria: list) -> tuple:
    """Return (score, summary_bullets) based on matching criteria."""
    score = 0
    summary = []
    text = fields.get('raw_text', '').lower()

    # Boolean/must-have criteria
    for crit in criteria:
        crit_low = crit.lower()
        if crit_low in text:
            score += 1
            summary.append(f"Matched: {crit}")

    # Skill-based weighting
    for skill in fields.get('skills', []):
        norm = normalize_skill(skill)
        if norm.lower() in [c.lower() for c in criteria]:
            score += 0.5
            summary.append(f"Skill: {norm}")

    return score, summary

